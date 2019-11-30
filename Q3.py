# CMT114 Coursework
# student number:

import docx
import os, sys
import time

def contain_numbers(text):
    list_text = text[1:-1].split(' ')
    if len(list_text) == 1:
        return False
    for val in list_text:
        if val.isdigit() and len(val)==4:
            return True
    return False


def change_style(filepath, style='IEEE'):
    if style == 'IEEE':
        f = open(filepath, 'rb')
        document = docx.Document(f)
        f.close()
        full_text = []
        for p in document.paragraphs:
            full_text.append(p.text)
        # print(full_text)
        replaced_text = []
        replaced_text = list(full_text)
        # print(replaced_text)

        replaceable_text = []
        for i, text in enumerate(full_text):
            if '(' in text and ')' in text:
                curl_index_open = 0
                curl_index_close = 0 
                while(1):
                    curl_index_open = text.find('(', curl_index_close, -1)
                    curl_index_close = text.find(')', curl_index_open, -1)
                    if curl_index_close == -1 or curl_index_open == -1:
                        break
                    else:
                        temp_text = text[curl_index_open:curl_index_close+1]
                        if contain_numbers(temp_text):
                            replaceable_text.append(temp_text)

        changed_replaceable_text = []
        for new_r_val in replaceable_text:
            new_r_val = new_r_val.replace('et al.', '').replace('e.g., ', '').replace('(', '').replace(')', '')
            new_r_val = new_r_val.split('; ')
            for string_val in new_r_val:
                if string_val not in changed_replaceable_text:
                    changed_replaceable_text.append(string_val)

        for i, ival in enumerate(replaced_text):
            for val in replaceable_text:
                if val in ival:
                    temp_text = []
                    for i_text, c_text in enumerate(changed_replaceable_text):
                        tt_val = val.replace('et al.', '').replace('e.g., ', '')
                        if c_text in val or c_text in tt_val:
                            temp_text.append([i_text+1])
                    replaced_text[i] = replaced_text[i].replace(val, '{}'.format(str(temp_text)[1:-1]))


        reference = []
        reference_start = False
        for val in replaced_text:
            # print(val)
            if 'Reference' in val:
                reference_start = True
            if reference_start:
                reference.append(val)

        new_reference = []
        for val in changed_replaceable_text:
            v_split = val.replace('and','').replace(',', '').split()
            for rf in reference:
                if v_split[0].lower() in rf.lower() and v_split[1].lower() in rf.lower():
                    new_reference.append(rf)
        
        document = docx.Document()
        references_text_start = False
        for i, text in enumerate(replaced_text):
            if 'Reference' in text :
                break
            else:
                document.add_paragraph(text)

        document.add_paragraph('References')
        print(dir(document))
        for i, text in enumerate(new_reference):
            document.add_paragraph( str(i+1) + '. ' + text)
        document.save('myfile_IEEE_style.docx')

    # APA style
    elif style == "APA":
        f = open(filepath, 'rb')
        document = docx.Document(f)
        f.close()

        full_text = []
        reference = []
        reference_style = []
        reference_start = False
        for p in document.paragraphs:
            reference_style.append(p)
            full_text.append(p.text)
            if reference_start :
                reference.append(p.text)
            if 'Reference' in p.text:
                reference_start = True

        replaced_text = []
        replaced_text = list(full_text)
        # print(replaced_text)

        new_reference = []
        reference_index = []
        for rf in reference:
            temp = rf[3:]
            if len(temp) > 10:
                t_index = temp.find(',')
                n_temp = temp[:t_index] + ', ' + temp[-6:-1]
                reference_index.append(n_temp)
                new_reference.append(temp)
        new_reference.sort()

        # print(replaced_text)
        # print("reference_index")
        # print(reference_index)

        for i, temp_rf_i in enumerate(reference_index):
            temp_text = '[{}]'.format(str(i))
            for index, rt in enumerate(replaced_text):
                if temp_text in rt:
                    replaced_text[index] = replaced_text[index].replace(temp_text, '({})'.format(str(temp_rf_i)))

        document = docx.Document()
        references_text_start = False
        for i, text in enumerate(replaced_text):
            if 'Reference' in text :
                break
            else:
                document.add_paragraph(text)

        document.add_paragraph('References')
        for i, text in enumerate(new_reference):
            document.add_paragraph(text)
        document.save('myfile_APA_style.docx')
            


# ---- DO NOT CHANGE THE CODE BELOW ----
if __name__ == "__main__":
    if len(sys.argv)<3: raise ValueError('Provide filename and style as input arguments')
    filepath, style = sys.argv[1], sys.argv[2]
    print('filepath is "{}"'.format(filepath))
    print('target style is "{}"'.format(style))
    change_style(filepath, style)
